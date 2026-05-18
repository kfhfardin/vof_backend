"""Word .docx extractor via python-docx."""

import io
from typing import BinaryIO, ClassVar

from docx import Document

from app.services.intake_extractors.base import ExtractedContent, IntakeExtractor


class DocxExtractor(IntakeExtractor):
    name: ClassVar[str] = "docx"
    accepts_mime: ClassVar[list[str]] = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    accepts_ext: ClassVar[list[str]] = [".docx"]

    def extract(self, blob: BinaryIO, filename: str) -> ExtractedContent:
        doc = Document(io.BytesIO(blob.read()))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        text = "\n\n".join(paragraphs)
        tables_data: list[list[list[str]]] = []
        for table in doc.tables:
            table_rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables_data.append(table_rows)
        return ExtractedContent(
            text=text or None,
            tables=tables_data or None,
            metadata={"paragraph_count": len(paragraphs), "table_count": len(tables_data)},
        )
