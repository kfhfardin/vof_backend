"""Unit tests for intake extractors + registry.

Each extractor runs against a fixture-shaped BytesIO; no network or DB.
"""

import io
import json

import pytest

from app.services.intake_extractors import UnsupportedFormat, registry
from app.services.intake_extractors.csv_extractor import CsvExtractor
from app.services.intake_extractors.docx_extractor import DocxExtractor
from app.services.intake_extractors.json_extractor import JsonExtractor
from app.services.intake_extractors.pdf_extractor import PdfExtractor
from app.services.intake_extractors.text_extractor import TextExtractor
from app.services.intake_extractors.xlsx_extractor import XlsxExtractor


def test_registry_has_all_seven_formats() -> None:
    names = set(registry.list_names())
    assert names == {"pdf", "docx", "text", "csv", "xlsx", "json"}
    # text covers both .txt and .md


def test_registry_resolves_by_mime() -> None:
    e = registry.resolve("application/json", "anything")
    assert e.name == "json"


def test_registry_resolves_by_extension_fallback() -> None:
    e = registry.resolve(None, "x.csv")
    assert e.name == "csv"


def test_registry_raises_for_unknown_format() -> None:
    with pytest.raises(UnsupportedFormat):
        registry.resolve("application/x-tarball", "x.tar.gz")


def test_text_extractor() -> None:
    e = TextExtractor()
    out = e.extract(io.BytesIO(b"hello world"), "note.txt")
    assert out.text == "hello world"
    assert out.rows is None


def test_text_extractor_handles_latin1_fallback() -> None:
    e = TextExtractor()
    # bytes invalid as UTF-8 but valid as Latin-1
    out = e.extract(io.BytesIO(b"caf\xe9"), "note.txt")
    assert out.text is not None
    assert "caf" in out.text


def test_json_extractor_array_of_objects_becomes_rows() -> None:
    e = JsonExtractor()
    data = [{"a": 1}, {"a": 2, "b": "x"}]
    out = e.extract(io.BytesIO(json.dumps(data).encode()), "x.json")
    assert out.rows == data
    assert out.metadata["row_count"] == 2


def test_json_extractor_object_becomes_text() -> None:
    e = JsonExtractor()
    out = e.extract(io.BytesIO(b'{"k":"v"}'), "x.json")
    assert out.rows is None
    assert out.text is not None and '"k"' in out.text


def test_json_extractor_invalid_yields_warning() -> None:
    e = JsonExtractor()
    out = e.extract(io.BytesIO(b"not json"), "x.json")
    assert out.warnings and "json_parse_failed" in out.warnings[0]


def test_csv_extractor() -> None:
    e = CsvExtractor()
    out = e.extract(io.BytesIO(b"name,role\nAlice,manager\nBob,rep\n"), "x.csv")
    assert out.rows == [
        {"name": "Alice", "role": "manager"},
        {"name": "Bob", "role": "rep"},
    ]
    assert out.metadata["row_count"] == 2


def test_csv_extractor_handles_utf8_bom() -> None:
    e = CsvExtractor()
    body = "﻿name,role\nAlice,manager\n".encode()
    out = e.extract(io.BytesIO(body), "x.csv")
    assert out.rows is not None and out.rows[0]["name"] == "Alice"


def test_xlsx_extractor_roundtrip() -> None:
    """Build a tiny xlsx in-memory and extract it."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.append(["name", "role"])
    ws.append(["Alice", "manager"])
    ws.append(["Bob", "rep"])
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)

    e = XlsxExtractor()
    out = e.extract(bio, "x.xlsx")
    assert out.rows is not None
    assert len(out.rows) == 2
    assert out.rows[0]["name"] == "Alice"
    assert out.metadata["row_count"] == 2


def test_docx_extractor_roundtrip() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Quarterly objection patterns")
    doc.add_paragraph("Buyers ask about integration timeline.")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    e = DocxExtractor()
    out = e.extract(bio, "playbook.docx")
    assert out.text is not None
    assert "objection patterns" in out.text
    assert out.metadata["paragraph_count"] == 2


def test_pdf_extractor_minimal_pdf() -> None:
    """Skip if reportlab isn't available; otherwise build a real 1-page PDF."""
    pytest.importorskip("reportlab", reason="reportlab needed to synthesize a test PDF")
    from reportlab.pdfgen import canvas

    bio = io.BytesIO()
    c = canvas.Canvas(bio)
    c.drawString(100, 750, "Hello from PDF")
    c.save()
    bio.seek(0)
    e = PdfExtractor()
    out = e.extract(bio, "x.pdf")
    assert out.metadata["page_count"] == 1
    # text extraction sometimes whitespace-pads; just check substring
    assert out.text is not None
    assert "Hello" in out.text


def test_pdf_extractor_scanned_only_warns() -> None:
    """A non-PDF byte stream should fail-soft with a warning, not crash."""
    e = PdfExtractor()
    out = e.extract(io.BytesIO(b"not a pdf"), "x.pdf")
    assert out.text is None
    assert any("failed" in w or "no_text_extracted" in w for w in out.warnings)
